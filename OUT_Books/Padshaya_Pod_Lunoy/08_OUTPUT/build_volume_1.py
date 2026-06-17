#!/usr/bin/env python3
"""Super Build — Падшая под Луной · Том 1 → .md + .docx + .txt"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
CHAPTERS_DIR = ROOT / "08_OUTPUT" / "Chapters"
OUT_DIR = ROOT / "08_OUTPUT" / "Book"

CHAPTER_ORDER = [f"CH_{i:03d}.md" for i in range(1, 9)]

TITLE_RU = "Падшая под Луной"
TITLE_EN = "Fallen Under the Moon"
SUBTITLE = "Том 1"
TAGLINE = "Она училась быть идеальной. Он научил её хотеть."
GENRE = "18+ · dark fantasy · dark romance"
HERO = "Элирия · POV 1 лицо · полуэльфийская принцесса"
BUILD_DATE = date.today().isoformat()
ENGINE = "OUT Prometheus v3.3"

ENGINE_SECTIONS = re.compile(
    r"^##\s+(CHAPTER_VECTOR|STATE_SNAPSHOT|Текст)\s*$", re.I
)
SCENE_HEADING = re.compile(r"^###\s+.+$")
YAML_FENCE = re.compile(r"^```")
END_CHAPTER = re.compile(r"^\*Конец главы", re.I)
UNIT_GATE = re.compile(r"^\*\*Unit Gate", re.I)
HORIZONTAL_RULE = re.compile(r"^---\s*$")
SCENE_BREAK = "* * *"


def extract_title_from_header(line: str) -> str:
    m = re.match(r"^#\s+CH_\d+\s*[—–-]\s*(.+)$", line.strip())
    return m.group(1).strip() if m else line.lstrip("# ").strip()


def parse_chapter_file(path: Path) -> tuple[int, str, list[str]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    chapter_index = int(re.search(r"CH_(\d+)", path.name).group(1))
    chapter_title = ""
    body_lines: list[str] = []
    in_text = False
    in_yaml = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("# ") and not chapter_title:
            chapter_title = extract_title_from_header(stripped)
            continue

        if ENGINE_SECTIONS.match(stripped):
            in_text = stripped.lower().endswith("текст")
            continue

        if YAML_FENCE.match(stripped):
            in_yaml = not in_yaml
            continue

        if in_yaml:
            continue

        if not in_text:
            continue

        if END_CHAPTER.match(stripped) or UNIT_GATE.match(stripped):
            break

        if stripped.startswith("## "):
            break

        if HORIZONTAL_RULE.match(stripped):
            continue

        if SCENE_HEADING.match(stripped):
            if body_lines and body_lines[-1].strip() != "":
                body_lines.append("")
            body_lines.append(SCENE_BREAK)
            body_lines.append("")
            continue

        body_lines.append(line.rstrip())

    paragraphs: list[str] = []
    buf: list[str] = []
    for line in body_lines:
        if line.strip() == "":
            if buf:
                paragraphs.append("\n".join(buf).strip())
                buf = []
        else:
            buf.append(line)
    if buf:
        paragraphs.append("\n".join(buf).strip())

    paragraphs = [p for p in paragraphs if p]
    while paragraphs and paragraphs[0] == SCENE_BREAK:
        paragraphs.pop(0)

    return chapter_index, chapter_title, paragraphs


def count_words(paragraphs: list[str]) -> int:
    text = " ".join(p.replace(SCENE_BREAK, "") for p in paragraphs)
    text = re.sub(r"\*+", "", text)
    return len(text.split())


def split_runs(text: str) -> list[tuple[str, bool, bool]]:
    parts: list[tuple[str, bool, bool]] = []
    pattern = re.compile(r"\*\*([^*]+)\*\*|\*([^*]+)\*")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last : m.start()], False, False))
        if m.group(1) is not None:
            parts.append((m.group(1), False, True))
        else:
            parts.append((m.group(2), True, False))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    if not parts:
        parts.append((text, False, False))
    return parts


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_end)


def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.35
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.0)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, italic: bool = False, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_body_paragraph(doc: Document, text: str) -> None:
    if text == SCENE_BREAK:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run("* * *")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.0)
    for segment, italic, bold in split_runs(text):
        if not segment:
            continue
        run = p.add_run(segment)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run.italic = italic
        run.bold = bold


def add_title_pages(doc: Document) -> None:
    add_centered(doc, TITLE_RU, 28, bold=True, space_after=12)
    add_centered(doc, TITLE_EN, 14, italic=True, space_after=18)
    add_centered(doc, SUBTITLE, 18, space_after=24)
    add_centered(doc, TAGLINE, 12, italic=True, space_after=18)
    add_centered(doc, GENRE, 11, italic=True, space_after=8)
    add_centered(doc, HERO, 11, space_after=24)
    add_centered(doc, f"{BUILD_DATE} · {ENGINE}", 10, space_after=0)
    doc.add_page_break()


def add_toc_page(doc: Document, units: list[tuple[int, str, list[str]]]) -> None:
    h = doc.add_heading("Оглавление", level=1)
    for run in h.runs:
        run.font.name = "Times New Roman"
    h.paragraph_format.first_line_indent = Cm(0)

    for idx, title, _ in units:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"Глава {idx}. {title}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_page_break()


def add_colophon(doc: Document, total_words: int, total_chars: int) -> None:
    doc.add_page_break()
    add_centered(doc, "— Конец тома 1 —", 14, bold=True, space_after=24)
    add_centered(doc, TAGLINE, 11, italic=True, space_after=18)
    add_centered(doc, f"Слов: ~{total_words:,}".replace(",", " "), 10, space_after=4)
    add_centered(doc, f"Знаков: {total_chars:,}".replace(",", " "), 10, space_after=12)
    add_centered(doc, f"Super Build · {BUILD_DATE}", 9, italic=True)


def setup_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{TITLE_RU} · {SUBTITLE} · ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    add_page_number_field(p)


def build_markdown(units: list[tuple[int, str, list[str]]], total_words: int, total_chars: int) -> str:
    out: list[str] = [
        "---",
        f'title: "{TITLE_RU}"',
        f'subtitle: "{SUBTITLE}"',
        f'title_en: "{TITLE_EN}"',
        f"date: {BUILD_DATE}",
        f"words: {total_words}",
        f"chars: {total_chars}",
        f"chapters: {len(units)}",
        f"engine: {ENGINE}",
        "---",
        "",
        f"# {TITLE_RU}",
        f"## {TITLE_EN}",
        f"### {SUBTITLE}",
        "",
        f"*{TAGLINE}*",
        "",
        f"**{GENRE}** · {HERO}",
        "",
        f"*Super Build: {BUILD_DATE} · {ENGINE}*",
        "",
        "---",
        "",
        "## Оглавление",
        "",
    ]

    for idx, title, _ in units:
        anchor = f"glava-{idx}"
        out.append(f"{idx}. [{title}](#{anchor})")

    out.extend(["", "---", ""])

    for i, (idx, title, paragraphs) in enumerate(units):
        out.append(f'<a id="glava-{idx}"></a>')
        out.append(f"## Глава {idx}. {title}")
        out.append("")
        for p in paragraphs:
            out.append(p)
            out.append("")
        if i < len(units) - 1:
            out.append("---")
            out.append("")

    out.extend([
        "---",
        "",
        "*— Конец тома 1 —*",
        "",
        f"*~{total_words} слов · {total_chars} знаков*",
        "",
    ])
    return "\n".join(out).strip() + "\n"


def build_plain_text(units: list[tuple[int, str, list[str]]]) -> str:
    out: list[str] = [
        TITLE_RU.upper(),
        TITLE_EN,
        SUBTITLE,
        "",
        TAGLINE,
        "",
        "=" * 60,
        "",
    ]
    for idx, title, paragraphs in units:
        out.append(f"ГЛАВА {idx}. {title.upper()}")
        out.append("")
        for p in paragraphs:
            if p == SCENE_BREAK:
                out.append("  * * *")
                out.append("")
            else:
                clean = re.sub(r"\*+", "", p)
                out.append(clean)
                out.append("")
        out.append("-" * 40)
        out.append("")
    out.append("— Конец тома 1 —")
    return "\n".join(out).strip() + "\n"


def build_docx(
    units: list[tuple[int, str, list[str]]],
    out_path: Path,
    total_words: int,
    total_chars: int,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    set_document_defaults(doc)
    setup_footer(doc)

    add_title_pages(doc)
    add_toc_page(doc, units)

    for i, (idx, title, paragraphs) in enumerate(units):
        if i > 0:
            doc.add_page_break()

        h = doc.add_heading(f"Глава {idx}. {title}", level=1)
        h.paragraph_format.first_line_indent = Cm(0)
        for run in h.runs:
            run.font.name = "Times New Roman"

        for para_text in paragraphs:
            add_body_paragraph(doc, para_text)

    add_colophon(doc, total_words, total_chars)
    doc.save(out_path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    units: list[tuple[int, str, list[str]]] = []
    char_counts: list[int] = []
    word_counts: list[int] = []

    for fname in CHAPTER_ORDER:
        path = CHAPTERS_DIR / fname
        if not path.exists():
            raise FileNotFoundError(path)
        idx, title, paragraphs = parse_chapter_file(path)
        units.append((idx, title, paragraphs))
        text = "\n\n".join(paragraphs)
        char_counts.append(len(text))
        word_counts.append(count_words(paragraphs))

    total_chars = sum(char_counts)
    total_words = sum(word_counts)

    slug = "Padshaya_pod_Lunoy_TOM_1"
    md_path = OUT_DIR / f"{slug}.md"
    docx_path = OUT_DIR / f"{slug}.docx"
    txt_path = OUT_DIR / f"{slug}.txt"

    md_path.write_text(build_markdown(units, total_words, total_chars), encoding="utf-8")
    txt_path.write_text(build_plain_text(units), encoding="utf-8")
    build_docx(units, docx_path, total_words, total_chars)

    manifest = f"""# Super Build manifest — Падшая под Луной · Том 1

**Дата:** {BUILD_DATE}  
**Движок:** {ENGINE}  
**Тип:** Super Build (титул · оглавление · колофон · нумерация · .txt)

## Объём

| Метрика | Значение |
|---------|----------|
| Глав | {len(CHAPTER_ORDER)} |
| Слов (проза) | ~{total_words:,} |
| Знаков (проза) | {total_chars:,} |

## Главы

| № | Название | Слов | Знаков |
|---|----------|------|--------|
"""
    for (idx, title, _), words, chars in zip(units, word_counts, char_counts):
        manifest += f"| {idx} | {title} | ~{words} | {chars} |\n"

    manifest += f"""
## Файлы

| Файл | Назначение |
|------|------------|
| `{md_path.name}` | полная рукопись + YAML + оглавление |
| `{docx_path.name}` | издание: титул, TOC, абзацный отступ, номера страниц |
| `{txt_path.name}` | plain text для читалок |
| `BUILD_MANIFEST.md` | этот отчёт |

## Источник

`08_OUTPUT/Chapters/CH_001.md` … `CH_008.md` — секция `## Текст`

**Исключено из сборки:** CHAPTER_VECTOR, STATE_SNAPSHOT, yaml, Unit Gate, engine `---`
"""
    (OUT_DIR / "BUILD_MANIFEST.md").write_text(manifest, encoding="utf-8")

    print("=== SUPER BUILD OK ===")
    print(f"MD:   {md_path}")
    print(f"DOCX: {docx_path}")
    print(f"TXT:  {txt_path}")
    print(f"Chapters: {len(units)} | words: ~{total_words} | chars: {total_chars}")


if __name__ == "__main__":
    main()