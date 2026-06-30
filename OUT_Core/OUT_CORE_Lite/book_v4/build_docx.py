import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "poluimya_korzina_yablok_v4.md"
OUT = ROOT / "poluimya_korzina_yablok_v4.docx"


def set_font(run, name="Georgia", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_runs_from_markdown(paragraph, text, base_size=11.5):
    # Handles simple **bold** and *italic* spans without trying to be a full markdown parser.
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_font(run, size=base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=base_size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=base_size, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=base_size)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
    normal.font.size = Pt(11.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, before, after, color in [
        ("Heading 1", 20, 18, 10, (85, 24, 31)),
        ("Heading 2", 15, 14, 8, (85, 24, 31)),
    ]:
        style = styles[style_name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    run = p.add_run("Полуимя")
    set_font(run, size=34, bold=True, color=(85, 24, 31))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Корзина яблок")
    set_font(run, size=20, italic=True, color=(46, 57, 49))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Dark anime romantasy")
    set_font(run, size=12, color=(90, 90, 90))

    doc.add_page_break()


def add_simple_toc(doc, chapters):
    p = doc.add_paragraph("Оглавление", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for title in chapters:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(title)
        set_font(run, size=11.5)
    doc.add_page_break()


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Полуимя: Корзина яблок")
    set_font(run, size=8.5, color=(110, 110, 110))


def build():
    md = SOURCE.read_text(encoding="utf-8")
    # Strip separators; chapter boundaries are handled by headings.
    raw_lines = [line.rstrip() for line in md.splitlines()]
    lines = []
    for line in raw_lines:
        if line.strip() == "---":
            continue
        lines.append(line)

    chapter_titles = [
        line.replace("# ", "").strip()
        for line in lines
        if re.match(r"^# Глава \d+\.", line)
    ]

    doc = Document()
    configure_document(doc)
    add_footer(doc.sections[0])
    add_title_page(doc)
    add_simple_toc(doc, chapter_titles)

    in_front_matter = True
    previous_blank = False
    first_chapter = True

    for line in lines:
        stripped = line.strip()
        if not stripped:
            previous_blank = True
            continue

        # Skip markdown front matter title/subtitle/toc from source; we create Word versions.
        if in_front_matter:
            if re.match(r"^# Глава \d+\.", stripped):
                in_front_matter = False
            else:
                continue

        if re.match(r"^# Глава \d+\.", stripped):
            if not first_chapter:
                doc.add_page_break()
            first_chapter = False
            p = doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            previous_blank = False
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
            previous_blank = False
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            add_runs_from_markdown(p, "• " + stripped[2:].strip())
            previous_blank = False
            continue

        p = doc.add_paragraph(style="Normal")
        # Dialogue lines read better without first-line indentation.
        if stripped.startswith(("—", "- ")):
            p.paragraph_format.first_line_indent = Inches(0)
        else:
            p.paragraph_format.first_line_indent = Inches(0.22)
        add_runs_from_markdown(p, stripped)
        previous_blank = False

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
